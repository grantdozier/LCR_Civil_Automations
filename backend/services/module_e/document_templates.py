"""
Module E - Document Templates
Template management and branding assets for consistent document generation
"""
from typing import Dict, Optional, List
from pathlib import Path
from dataclasses import dataclass
import logging

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


@dataclass
class BrandingAssets:
    """
    Branding assets for document generation.

    Attributes:
        company_name: Full company name
        tagline: Company tagline/slogan
        address: Mailing address
        phone: Phone number
        email: Contact email
        website: Company website
        logo_path: Path to company logo (optional)
        primary_color: Primary brand color (RGB tuple)
        secondary_color: Secondary brand color (RGB tuple)
        letterhead_template: Path to letterhead template (optional)
    """
    company_name: str
    tagline: str
    address: str
    phone: str
    email: str
    website: str
    logo_path: Optional[str] = None
    primary_color: tuple = (0, 51, 102)  # Dark blue
    secondary_color: tuple = (0, 128, 0)  # Green
    letterhead_template: Optional[str] = None


class TemplateManager:
    """
    Manage document templates and branding for consistent output.

    Provides:
    - Standard templates for common documents
    - Branding application (logos, colors, fonts)
    - Template customization
    - Reusable document components
    """

    # Standard branding assets
    LCR_BRANDING = BrandingAssets(
        company_name="LCR & Company",
        tagline="Civil Engineering & Land Surveying",
        address="Lafayette, Louisiana 70508",
        phone="(337) 555-1234",
        email="info@lcrcompany.com",
        website="www.lcrcompany.com",
        primary_color=(0, 51, 102),
        secondary_color=(0, 128, 0),
    )

    DOZIER_BRANDING = BrandingAssets(
        company_name="Dozier Tech",
        tagline="Engineering Technology Solutions",
        address="Lafayette, Louisiana 70508",
        phone="(337) 555-5678",
        email="grant@doziertech.com",
        website="www.doziertech.com",
        primary_color=(51, 51, 51),
        secondary_color=(0, 102, 204),
    )

    def __init__(self, templates_dir: str = "/app/templates"):
        """
        Initialize template manager.

        Args:
            templates_dir: Directory containing template files
        """
        self.templates_dir = Path(templates_dir)
        self.templates_dir.mkdir(parents=True, exist_ok=True)

    def get_branding(self, company: str) -> BrandingAssets:
        """
        Get branding assets for specified company.

        Args:
            company: "LCR", "Dozier", or "Both"

        Returns:
            BrandingAssets object
        """
        if company.upper() == "LCR":
            return self.LCR_BRANDING
        elif company.upper() == "DOZIER":
            return self.DOZIER_BRANDING
        else:
            # Combined branding
            return BrandingAssets(
                company_name="LCR & Company | Dozier Tech",
                tagline="Integrated Engineering Solutions",
                address="Lafayette, Louisiana 70508",
                phone="(337) 555-1234",
                email="info@lcrcompany.com",
                website="www.lcrcompany.com",
                primary_color=(0, 51, 102),
                secondary_color=(0, 128, 0),
            )

    def create_branded_document(
        self,
        company: str,
        document_type: str = "letter"
    ) -> Document:
        """
        Create a new document with company branding.

        Args:
            company: "LCR", "Dozier", or "Both"
            document_type: Type of document ("letter", "report", "proposal")

        Returns:
            Document object with branding applied
        """
        branding = self.get_branding(company)
        doc = Document()

        # Apply branding based on document type
        if document_type == "letter":
            self._add_letterhead(doc, branding)
        elif document_type == "report":
            self._add_report_header(doc, branding)
        elif document_type == "proposal":
            self._add_proposal_cover(doc, branding)

        return doc

    def _add_letterhead(self, doc: Document, branding: BrandingAssets):
        """Add company letterhead to document"""
        # Header with company name
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company name
        run = header.add_run(f"{branding.company_name}\\n")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*branding.primary_color)

        # Tagline
        run = header.add_run(f"{branding.tagline}\\n")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*branding.primary_color)

        # Contact info
        contact_info = f"{branding.address} | {branding.phone} | {branding.email}"
        run = header.add_run(contact_info)
        run.font.size = Pt(9)

        doc.add_paragraph("")  # Spacing

    def _add_report_header(self, doc: Document, branding: BrandingAssets):
        """Add report header to document"""
        # Logo/Company name
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = header.add_run(f"{branding.company_name}\\n")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*branding.primary_color)

        run = header.add_run(f"{branding.tagline}\\n\\n")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*branding.primary_color)

    def _add_proposal_cover(self, doc: Document, branding: BrandingAssets):
        """Add proposal cover page to document"""
        self._add_report_header(doc, branding)

    def apply_standard_styles(self, doc: Document):
        """
        Apply standard formatting styles to document.

        Args:
            doc: Document to apply styles to
        """
        # This would set standard fonts, spacing, etc.
        # For now, just a placeholder
        logger.info("Standard styles applied")

    def get_template_path(self, template_name: str) -> Optional[Path]:
        """
        Get path to a template file.

        Args:
            template_name: Name of template (e.g., "proposal_template.docx")

        Returns:
            Path to template file if exists, None otherwise
        """
        template_path = self.templates_dir / template_name
        if template_path.exists():
            return template_path
        else:
            logger.warning(f"Template not found: {template_path}")
            return None

    def create_template_from_document(
        self,
        doc: Document,
        template_name: str
    ) -> str:
        """
        Save a document as a reusable template.

        Args:
            doc: Document to save as template
            template_name: Name for the template file

        Returns:
            Path to saved template
        """
        template_path = self.templates_dir / template_name
        doc.save(str(template_path))

        logger.info(f"Created template: {template_path}")
        return str(template_path)

    def get_standard_terms_and_conditions(self, service_type: str = "general") -> List[str]:
        """
        Get standard terms and conditions text.

        Args:
            service_type: Type of service ("general", "automation", "survey")

        Returns:
            List of terms and conditions paragraphs
        """
        if service_type == "automation":
            return [
                "Scope of Work: Services are limited to the automation modules specified in this agreement.",
                "Client Data: Client is responsible for providing accurate and complete project data in a timely manner.",
                "Deliverables: All deliverables will be provided in industry-standard digital formats (Word, PDF, DWG, etc.).",
                "Revisions: One round of revisions is included. Additional revisions will be billed at our standard hourly rate.",
                "Payment Terms: Payment is due net 30 days from invoice date unless otherwise specified.",
                "Intellectual Property: Client retains ownership of all deliverables upon final payment.",
                "Warranty: We warrant that all work will be performed in accordance with accepted civil engineering practices.",
                "Limitation of Liability: Our liability is limited to the fee paid for services.",
                "Force Majeure: We are not liable for delays due to circumstances beyond our reasonable control.",
                "Governing Law: This agreement is governed by the laws of the State of Louisiana.",
            ]
        elif service_type == "survey":
            return [
                "Scope of Survey: Survey services are limited to the property boundaries and features specified.",
                "Access: Client will provide access to property and obtain necessary permissions.",
                "Field Conditions: Survey is based on observable field conditions at time of survey.",
                "Monumentation: Property corners will be marked with iron rods or monuments as required.",
                "Professional Standards: All work performed in accordance with Louisiana R.S. 37:682-37:704.",
            ]
        else:
            return [
                "Scope: Services are limited to those specifically described in this agreement.",
                "Standard of Care: Services performed in accordance with accepted professional engineering practices.",
                "Payment: Invoices are due net 30 days from date of invoice.",
                "Changes: Changes to scope will be documented via written change order.",
                "Ownership: Client owns all deliverables upon final payment.",
            ]

    def get_standard_disclaimers(self) -> List[str]:
        """
        Get standard professional disclaimers.

        Returns:
            List of disclaimer paragraphs
        """
        return [
            "This document has been prepared by or under the supervision of a Louisiana Licensed Professional Engineer.",
            "This analysis is based on information provided by the client and publicly available data sources.",
            "Field verification is recommended prior to final design and construction.",
            "Subsurface conditions may vary from those assumed. Geotechnical investigation is recommended.",
            "Rainfall data based on NOAA Atlas 14, Volume 9 (Southeastern States).",
            "All calculations performed in accordance with accepted civil engineering practice.",
        ]

    def format_currency(self, amount: float) -> str:
        """
        Format currency value consistently.

        Args:
            amount: Dollar amount

        Returns:
            Formatted currency string
        """
        return f"${amount:,.2f}"

    def format_date(self, date: str = None) -> str:
        """
        Format date consistently.

        Args:
            date: Date string or None for today

        Returns:
            Formatted date string
        """
        from datetime import datetime
        if date:
            # Parse and reformat
            # For now, just return as-is
            return date
        else:
            return datetime.now().strftime("%B %d, %Y")

    def create_signature_block(
        self,
        doc: Document,
        name: str,
        title: Optional[str] = None,
        company: Optional[str] = None
    ):
        """
        Add a signature block to document.

        Args:
            doc: Document to add signature block to
            name: Name of person signing
            title: Job title or credentials (e.g., "PE", "PLS")
            company: Company name
        """
        doc.add_paragraph("")
        doc.add_paragraph("Respectfully submitted,")
        doc.add_paragraph("")
        doc.add_paragraph("")

        sig = doc.add_paragraph()
        sig.add_run(name).font.bold = True

        if title:
            doc.add_paragraph(title)

        if company:
            doc.add_paragraph(company)

    def create_project_info_table(
        self,
        doc: Document,
        project_data: Dict
    ):
        """
        Add a standard project information table.

        Args:
            doc: Document to add table to
            project_data: Dictionary with project information
        """
        # Standard fields
        fields = [
            ("Project Name:", project_data.get("project_name", "N/A")),
            ("Project Number:", project_data.get("project_number", "N/A")),
            ("Location:", project_data.get("location", "N/A")),
            ("Client:", project_data.get("client_name", "N/A")),
            ("Date:", project_data.get("date", self.format_date())),
        ]

        # Create table
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = 'Light Grid Accent 1'

        for idx, (label, value) in enumerate(fields):
            table.rows[idx].cells[0].text = label
            table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[idx].cells[1].text = value

        doc.add_paragraph("")  # Spacing


class DocumentBatcher:
    """
    Batch generate multiple documents with consistent formatting.

    Useful for generating multiple proposals, cover letters, or reports
    with different data but same template.
    """

    def __init__(self, template_manager: TemplateManager):
        """
        Initialize document batcher.

        Args:
            template_manager: TemplateManager instance
        """
        self.template_manager = template_manager

    def batch_generate_proposals(
        self,
        proposals_data: List[Dict],
        output_dir: Path
    ) -> List[str]:
        """
        Generate multiple proposals from list of data.

        Args:
            proposals_data: List of proposal data dictionaries
            output_dir: Directory for output files

        Returns:
            List of paths to generated files
        """
        from services.module_e.proposal_generator import (
            ProposalGenerator,
            ProposalData
        )

        output_paths = []
        generator = ProposalGenerator(output_dir=str(output_dir))

        for data in proposals_data:
            # Convert dict to ProposalData
            proposal = ProposalData(**data)

            # Generate proposal
            path = generator.generate_proposal(proposal)
            output_paths.append(path)

        logger.info(f"Batch generated {len(output_paths)} proposals")
        return output_paths

    def batch_generate_cover_letters(
        self,
        submittals_data: List[Dict],
        output_dir: Path
    ) -> List[str]:
        """
        Generate multiple cover letters from list of data.

        Args:
            submittals_data: List of submittal data dictionaries
            output_dir: Directory for output files

        Returns:
            List of paths to generated files
        """
        from services.module_e.cover_letter_generator import (
            CoverLetterGenerator,
            SubmittalPackage
        )

        output_paths = []
        generator = CoverLetterGenerator(output_dir=str(output_dir))

        for data in submittals_data:
            # Convert dict to SubmittalPackage
            submittal = SubmittalPackage(**data)

            # Generate cover letter
            path = generator.generate_cover_letter(submittal)
            output_paths.append(path)

        logger.info(f"Batch generated {len(output_paths)} cover letters")
        return output_paths
