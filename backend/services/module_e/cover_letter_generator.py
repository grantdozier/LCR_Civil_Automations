"""
Module E - Cover Letter Generator
Generate professional cover letters and submittal packages
"""
from typing import List, Dict, Optional
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime
import logging

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


@dataclass
class SubmittalDocument:
    """A document included in the submittal package"""
    document_type: str  # e.g., "DIA Report", "Drainage Plans", "QA Report"
    description: str
    filename: str
    page_count: Optional[int] = None
    revision: Optional[str] = None


@dataclass
class SubmittalPackage:
    """
    Complete submittal package data.

    Attributes:
        project_name: Name of the project
        project_number: Project number
        client_name: Client organization
        client_address: Mailing address
        client_contact: Primary contact person
        subject: Subject line for letter
        documents: List of documents being submitted
        purpose: Purpose of submittal (e.g., "Plan Review", "Final Approval")
        special_instructions: Any special notes or instructions
        prepared_by: Name of person preparing submittal
        company: "LCR" or "Dozier" or "Both"
    """
    project_name: str
    project_number: str
    client_name: str
    client_address: str
    client_contact: str
    subject: str
    documents: List[SubmittalDocument]
    purpose: str = "Plan Review"
    special_instructions: Optional[str] = None
    prepared_by: str = "Grant Dozier, PE"
    company: str = "LCR"


class CoverLetterGenerator:
    """
    Generate professional cover letters and submittal packages.

    Creates:
    - Cover letters for plan submittals
    - Document transmittal letters
    - Response letters to review comments
    - Professional correspondence with branding
    """

    def __init__(self, output_dir: str = "/app/outputs"):
        """
        Initialize cover letter generator.

        Args:
            output_dir: Directory for output files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_cover_letter(
        self,
        submittal: SubmittalPackage,
        letter_type: str = "submittal",
        output_filename: Optional[str] = None
    ) -> str:
        """
        Generate professional cover letter.

        Args:
            submittal: Submittal package data
            letter_type: Type of letter ("submittal", "transmittal", "response")
            output_filename: Optional custom filename

        Returns:
            Path to generated cover letter file
        """
        logger.info(f"Generating {letter_type} cover letter for {submittal.project_name}")

        # Create document
        doc = Document()

        # Letterhead
        self._add_letterhead(doc, submittal.company)

        # Date
        date_para = doc.add_paragraph()
        date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        doc.add_paragraph("")  # Spacing

        # Recipient address
        self._add_recipient_address(doc, submittal)
        doc.add_paragraph("")

        # Subject line
        subject_para = doc.add_paragraph()
        subject_para.add_run("RE: ").font.bold = True
        subject_para.add_run(submittal.subject)
        doc.add_paragraph("")

        # Salutation
        doc.add_paragraph(f"Dear {submittal.client_contact}:")
        doc.add_paragraph("")

        # Body
        if letter_type == "submittal":
            self._add_submittal_body(doc, submittal)
        elif letter_type == "transmittal":
            self._add_transmittal_body(doc, submittal)
        elif letter_type == "response":
            self._add_response_body(doc, submittal)
        else:
            self._add_generic_body(doc, submittal)

        # Document list
        self._add_document_list(doc, submittal)

        # Special instructions
        if submittal.special_instructions:
            doc.add_paragraph("")
            notes_para = doc.add_paragraph()
            notes_para.add_run("Special Notes:\\n").font.bold = True
            notes_para.add_run(submittal.special_instructions)

        # Closing
        self._add_closing(doc, submittal)

        # Save document
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d")
            safe_project = submittal.project_number.replace(" ", "_")
            output_filename = f"Cover_Letter_{safe_project}_{timestamp}.docx"

        output_path = self.output_dir / output_filename
        doc.save(str(output_path))

        logger.info(f"Generated cover letter: {output_path}")
        return str(output_path)

    def _add_letterhead(self, doc: Document, company: str):
        """Add company letterhead"""
        # Company name and logo
        letterhead = doc.add_paragraph()
        letterhead.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if company == "LCR":
            company_name = "LCR & COMPANY"
            tagline = "Civil Engineering & Land Surveying"
            contact = "Lafayette, Louisiana | (337) 555-1234 | info@lcrcompany.com"
        elif company == "Dozier":
            company_name = "DOZIER TECH"
            tagline = "Engineering Technology Solutions"
            contact = "Lafayette, Louisiana | (337) 555-5678 | grant@doziertech.com"
        else:
            company_name = "LCR & COMPANY | DOZIER TECH"
            tagline = "Integrated Engineering Solutions"
            contact = "Lafayette, Louisiana | (337) 555-1234"

        run = letterhead.add_run(f"{company_name}\\n")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        run = letterhead.add_run(f"{tagline}\\n")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 51, 102)

        run = letterhead.add_run(contact)
        run.font.size = Pt(9)

        doc.add_paragraph("")  # Spacing

    def _add_recipient_address(self, doc: Document, submittal: SubmittalPackage):
        """Add recipient address block"""
        address = doc.add_paragraph()
        address.add_run(f"{submittal.client_contact}\\n")
        address.add_run(f"{submittal.client_name}\\n")

        # Split address into lines
        for line in submittal.client_address.split("\\n"):
            address.add_run(f"{line.strip()}\\n")

    def _add_submittal_body(self, doc: Document, submittal: SubmittalPackage):
        """Add body text for plan submittal letter"""
        intro = doc.add_paragraph()
        intro.add_run(
            f"On behalf of {submittal.client_name}, we are pleased to submit the following "
            f"documents for {submittal.purpose.lower()} for the "
            f"{submittal.project_name} project."
        )

        doc.add_paragraph("")

        context = doc.add_paragraph()
        context.add_run(
            "This submittal package includes all required documentation prepared in accordance with "
            "Lafayette Unified Development Code (UDC), Louisiana Department of Transportation and "
            "Development (DOTD) standards, and Louisiana Pollutant Discharge Elimination System (LPDES) "
            "requirements."
        )

    def _add_transmittal_body(self, doc: Document, submittal: SubmittalPackage):
        """Add body text for document transmittal letter"""
        intro = doc.add_paragraph()
        intro.add_run(
            f"Enclosed please find the following documents for the "
            f"{submittal.project_name} project:"
        )

    def _add_response_body(self, doc: Document, submittal: SubmittalPackage):
        """Add body text for response to review comments"""
        intro = doc.add_paragraph()
        intro.add_run(
            f"Thank you for your review of the {submittal.project_name} project. "
            f"We have addressed all comments and are resubmitting the following "
            f"revised documents for your consideration:"
        )

    def _add_generic_body(self, doc: Document, submittal: SubmittalPackage):
        """Add generic body text"""
        intro = doc.add_paragraph()
        intro.add_run(
            f"Please find enclosed the following documents regarding the "
            f"{submittal.project_name} project:"
        )

    def _add_document_list(self, doc: Document, submittal: SubmittalPackage):
        """Add list of enclosed documents"""
        doc.add_paragraph("")
        doc.add_heading("Enclosed Documents:", level=3)

        # Create document table
        table = doc.add_table(rows=len(submittal.documents) + 1, cols=4)
        table.style = 'Light Grid Accent 1'

        # Headers
        headers = ["Item", "Document Type", "Description", "Pages/Revision"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Document rows
        for idx, doc_item in enumerate(submittal.documents, start=1):
            row = table.rows[idx]
            row.cells[0].text = str(idx)
            row.cells[1].text = doc_item.document_type
            row.cells[2].text = doc_item.description

            # Pages/revision
            info = []
            if doc_item.page_count:
                info.append(f"{doc_item.page_count} pages")
            if doc_item.revision:
                info.append(f"Rev. {doc_item.revision}")
            row.cells[3].text = ", ".join(info) if info else "-"

        doc.add_paragraph("")

        # Summary
        total_docs = len(submittal.documents)
        summary = doc.add_paragraph()
        summary.add_run(f"Total Documents Enclosed: {total_docs}")

    def _add_closing(self, doc: Document, submittal: SubmittalPackage):
        """Add closing and signature block"""
        doc.add_paragraph("")

        closing = doc.add_paragraph()
        closing.add_run(
            "Should you have any questions or require additional information, "
            "please do not hesitate to contact our office."
        )

        doc.add_paragraph("")
        doc.add_paragraph("Respectfully submitted,")
        doc.add_paragraph("")
        doc.add_paragraph("")

        # Signature block
        sig = doc.add_paragraph()
        sig.add_run(submittal.prepared_by).font.bold = True

        # Add title/credentials if in name
        if ", PE" in submittal.prepared_by or "P.E." in submittal.prepared_by:
            doc.add_paragraph("Professional Engineer")
        elif ", PLS" in submittal.prepared_by:
            doc.add_paragraph("Professional Land Surveyor")

        if submittal.company == "LCR":
            doc.add_paragraph("LCR & Company")
        elif submittal.company == "Dozier":
            doc.add_paragraph("Dozier Tech")
        else:
            doc.add_paragraph("LCR & Company | Dozier Tech")

    def generate_transmittal_form(
        self,
        submittal: SubmittalPackage,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Generate a formal document transmittal form.

        Args:
            submittal: Submittal package data
            output_filename: Optional custom filename

        Returns:
            Path to generated transmittal form
        """
        logger.info(f"Generating transmittal form for {submittal.project_name}")

        doc = Document()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("DOCUMENT TRANSMITTAL FORM\\n\\n")
        run.font.size = Pt(16)
        run.font.bold = True

        # Project info table
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'

        # Fill in project information
        info_data = [
            ("Project Name:", submittal.project_name),
            ("Project Number:", submittal.project_number),
            ("To:", submittal.client_name),
            ("Attention:", submittal.client_contact),
            ("From:", submittal.prepared_by),
            ("Date:", datetime.now().strftime("%B %d, %Y")),
        ]

        for idx, (label, value) in enumerate(info_data):
            info_table.rows[idx].cells[0].text = label
            info_table.rows[idx].cells[0].paragraphs[0].runs[0].font.bold = True
            info_table.rows[idx].cells[1].text = value

        doc.add_paragraph("")

        # Purpose
        purpose_para = doc.add_paragraph()
        purpose_para.add_run("Purpose:\\n").font.bold = True
        purpose_para.add_run(submittal.purpose)

        doc.add_paragraph("")

        # Documents transmitted
        doc.add_heading("Documents Transmitted:", level=2)

        # Checklist style
        for idx, doc_item in enumerate(submittal.documents, start=1):
            item_para = doc.add_paragraph()
            item_para.add_run(f"☐ {idx}. {doc_item.document_type}\\n")
            item_para.add_run(f"     {doc_item.description}")

            if doc_item.page_count or doc_item.revision:
                info = []
                if doc_item.page_count:
                    info.append(f"{doc_item.page_count} pages")
                if doc_item.revision:
                    info.append(f"Revision {doc_item.revision}")
                item_para.add_run(f" ({', '.join(info)})")

        doc.add_paragraph("")

        # Transmittal options
        doc.add_heading("Transmittal Action:", level=2)

        actions = [
            "☐ For Review and Comment",
            "☐ For Approval",
            "☐ For Information Only",
            "☐ As Requested",
            "☐ For Record",
        ]

        for action in actions:
            doc.add_paragraph(action)

        doc.add_paragraph("")

        # Remarks
        remarks = doc.add_paragraph()
        remarks.add_run("Remarks/Special Instructions:\\n").font.bold = True
        if submittal.special_instructions:
            remarks.add_run(submittal.special_instructions)
        else:
            remarks.add_run("\\n" * 3)

        doc.add_paragraph("")

        # Signature
        sig = doc.add_paragraph()
        sig.add_run("Submitted By: _" + "_" * 40 + "  Date: _" + "_" * 20)

        # Save
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d")
            safe_project = submittal.project_number.replace(" ", "_")
            output_filename = f"Transmittal_Form_{safe_project}_{timestamp}.docx"

        output_path = self.output_dir / output_filename
        doc.save(str(output_path))

        logger.info(f"Generated transmittal form: {output_path}")
        return str(output_path)

    def generate_review_response_letter(
        self,
        submittal: SubmittalPackage,
        review_comments: List[Dict[str, str]],
        output_filename: Optional[str] = None
    ) -> str:
        """
        Generate response letter to review comments.

        Args:
            submittal: Submittal package data
            review_comments: List of comments with responses
                             Each dict should have: {"comment": str, "response": str}
            output_filename: Optional custom filename

        Returns:
            Path to generated response letter
        """
        logger.info(f"Generating review response letter for {submittal.project_name}")

        doc = Document()

        # Letterhead
        self._add_letterhead(doc, submittal.company)

        # Date
        date_para = doc.add_paragraph()
        date_para.add_run(datetime.now().strftime("%B %d, %Y"))
        doc.add_paragraph("")

        # Recipient
        self._add_recipient_address(doc, submittal)
        doc.add_paragraph("")

        # Subject
        subject_para = doc.add_paragraph()
        subject_para.add_run("RE: ").font.bold = True
        subject_para.add_run(f"Response to Review Comments - {submittal.subject}")
        doc.add_paragraph("")

        # Salutation
        doc.add_paragraph(f"Dear {submittal.client_contact}:")
        doc.add_paragraph("")

        # Introduction
        intro = doc.add_paragraph()
        intro.add_run(
            f"Thank you for your review of the {submittal.project_name} project. "
            f"We have carefully considered all comments and have made the necessary "
            f"revisions. Below is our response to each comment:"
        )

        doc.add_paragraph("")

        # Comment/Response table
        doc.add_heading("Review Comments and Responses:", level=2)

        for idx, item in enumerate(review_comments, start=1):
            # Comment number
            comment_heading = doc.add_paragraph()
            comment_heading.add_run(f"Comment {idx}:\\n").font.bold = True
            comment_heading.add_run(item.get("comment", ""))

            # Response
            response_para = doc.add_paragraph()
            response_para.add_run("Response: ").font.bold = True
            response_para.add_run(item.get("response", ""))

            doc.add_paragraph("")  # Spacing

        # Closing
        doc.add_paragraph("")
        closing = doc.add_paragraph()
        closing.add_run(
            "We believe all comments have been adequately addressed. "
            "Please review the revised submittal and contact us with any questions."
        )

        self._add_closing(doc, submittal)

        # Save
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d")
            safe_project = submittal.project_number.replace(" ", "_")
            output_filename = f"Review_Response_{safe_project}_{timestamp}.docx"

        output_path = self.output_dir / output_filename
        doc.save(str(output_path))

        logger.info(f"Generated review response letter: {output_path}")
        return str(output_path)
