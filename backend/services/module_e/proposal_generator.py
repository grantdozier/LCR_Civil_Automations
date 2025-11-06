"""
Module E - Proposal Generator
Auto-generate branded proposals for civil engineering services
Generates client proposals, submittal letters, and fee summaries for LCR's drainage engineering services
"""
from typing import List, Dict, Optional
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from decimal import Decimal
import logging

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


@dataclass
class ServicePricing:
    """Pricing for individual civil engineering services"""
    service_id: str
    service_name: str
    description: str
    base_price: Decimal
    time_estimate_days: int
    deliverables: List[str] = field(default_factory=list)
    unit: str = "per project"  # "per project", "per acre", "per sheet", etc.


@dataclass
class ProposalData:
    """
    Data for generating a civil engineering services proposal.

    Attributes:
        client_name: Name of client organization (e.g., "Lafayette Parish School Board")
        client_contact: Primary contact person
        client_email: Contact email
        project_name: Name of the project (e.g., "L.J. Alleman Middle School Drainage Improvements")
        project_location: Project location (e.g., "Lafayette, LA")
        project_description: Brief description of the project
        jurisdiction: Regulatory jurisdiction (e.g., "Lafayette UDC", "DOTD", "LCG")
        project_type: Type of project (e.g., "Educational", "Commercial", "Residential", "Municipal")
        services_requested: List of service IDs (e.g., ["DIA", "GRADING", "REVIEW"])
        project_area_acres: Optional project area in acres (for area-based pricing)
        num_plan_sheets: Optional number of plan sheets (for sheet-based pricing)
        custom_scope: Optional custom scope items
        discount_percent: Optional discount percentage (0-100)
        rush_fee_percent: Optional rush fee percentage (0-100)
        proposal_valid_days: Days proposal is valid (default 30)
        prepared_by: Name of person preparing proposal
        company: "LCR" (LCR & Company)
    """
    client_name: str
    client_contact: str
    client_email: str
    project_name: str
    project_location: str
    project_description: str
    jurisdiction: str = "Lafayette UDC"
    project_type: str = "Commercial"
    services_requested: List[str] = field(default_factory=list)
    project_area_acres: Optional[Decimal] = None
    num_plan_sheets: Optional[int] = None
    custom_scope: List[str] = field(default_factory=list)
    discount_percent: Decimal = Decimal("0")
    rush_fee_percent: Decimal = Decimal("0")
    proposal_valid_days: int = 30
    prepared_by: str = "Grant Dozier, PE"
    company: str = "LCR"


class PricingCalculator:
    """
    Calculate pricing for civil engineering service combinations.

    Standard Civil Engineering Services offered by LCR & Company:
    - DIA: Drainage Impact Analysis - Complete drainage study with Rational Method
    - GRADING: Grading Plan Review & Approval Assistance
    - DETENTION: Detention Pond Design & Analysis
    - TOC: Time of Concentration Analysis
    - REVIEW: Plan Review & Compliance Checks (LPDES/LUS/DOTD)
    - SURVEY: Survey Coordination & Review
    - SUBMITTAL: Submittal Package Preparation
    - CONSTRUCTION: Construction Observation Services
    - STORMWATER: Stormwater Management Plan

    Package Discounts:
    - 3+ services: 5% discount
    - 5+ services: 10% discount
    - Full package (7+ services): 15% discount
    """

    # Standard civil engineering service pricing
    SERVICE_PRICING = {
        "DIA": ServicePricing(
            service_id="DIA",
            service_name="Drainage Impact Analysis (DIA Report)",
            description="Complete drainage study with Rational Method calculations, pre/post development analysis, and regulatory compliance documentation",
            base_price=Decimal("4500"),
            time_estimate_days=10,
            deliverables=[
                "58+ page professional DIA report",
                "Rational Method calculations (Q=CiA)",
                "Time of Concentration analysis (4 methods)",
                "Multi-storm event analysis (10/25/50/100-year)",
                "Pre and post-development comparison",
                "Technical exhibits and tables",
                "NOAA Atlas 14 rainfall data",
                "UDC/DOTD compliance documentation",
                "Client-ready PDF and Word documents",
            ],
            unit="per project"
        ),
        "GRADING": ServicePricing(
            service_id="GRADING",
            service_name="Grading Plan Review & Design",
            description="Professional grading plan development and review for drainage compliance",
            base_price=Decimal("3500"),
            time_estimate_days=8,
            deliverables=[
                "Grading plan review and markup",
                "Spot elevation verification",
                "Drainage flow path analysis",
                "Compliance with local standards",
                "Grading calculations and documentation",
                "Coordination with Civil 3D plans",
            ],
            unit="per project"
        ),
        "DETENTION": ServicePricing(
            service_id="DETENTION",
            service_name="Detention Pond Design & Analysis",
            description="Stormwater detention facility design with release rate calculations",
            base_price=Decimal("5000"),
            time_estimate_days=12,
            deliverables=[
                "Detention pond sizing calculations",
                "Hydraulic analysis and routing",
                "Outlet structure design",
                "Emergency spillway design",
                "Stage-storage-discharge curves",
                "Release rate compliance verification",
                "Construction details and specifications",
            ],
            unit="per facility"
        ),
        "TOC": ServicePricing(
            service_id="TOC",
            service_name="Time of Concentration Analysis",
            description="Professional TOC calculations using multiple approved methods",
            base_price=Decimal("1500"),
            time_estimate_days=3,
            deliverables=[
                "TOC calculations (Kerby, Kirpich, NRCS, Kinematic Wave)",
                "Flow path delineation",
                "Weighted C-value calculations",
                "Area-weighted TOC analysis",
                "Compliance verification",
                "Professional calculation sheets",
            ],
            unit="per basin"
        ),
        "REVIEW": ServicePricing(
            service_id="REVIEW",
            service_name="Plan Review & QA Services",
            description="Comprehensive plan review for regulatory compliance before submittal",
            base_price=Decimal("2500"),
            time_estimate_days=5,
            deliverables=[
                "Complete plan set review (C-1 through C-18)",
                "LPDES compliance checklist",
                "LUS/LCG/DOTD standard verification",
                "ASTM specification checks",
                "Standard notes verification",
                "Detail and callout review",
                "Professional QA report with redlines",
            ],
            unit="per submittal"
        ),
        "SURVEY": ServicePricing(
            service_id="SURVEY",
            service_name="Survey Coordination & Review",
            description="Survey data review and coordination for drainage design",
            base_price=Decimal("2000"),
            time_estimate_days=4,
            deliverables=[
                "Survey data review and validation",
                "Topographic data QA/QC",
                "Existing infrastructure verification",
                "Benchmark and control point review",
                "Coordinate system verification",
                "Survey data integration with Civil 3D",
            ],
            unit="per project"
        ),
        "SUBMITTAL": ServicePricing(
            service_id="SUBMITTAL",
            service_name="Submittal Package Preparation",
            description="Professional submittal package assembly with cover letters and transmittals",
            base_price=Decimal("1200"),
            time_estimate_days=2,
            deliverables=[
                "Professional cover letter",
                "Document transmittal forms",
                "Submittal checklist compliance",
                "Digital and print package assembly",
                "Agency-specific formatting",
                "Tracking and follow-up documentation",
            ],
            unit="per submittal"
        ),
        "CONSTRUCTION": ServicePricing(
            service_id="CONSTRUCTION",
            service_name="Construction Observation Services",
            description="On-site construction observation and compliance verification",
            base_price=Decimal("1800"),
            time_estimate_days=1,
            deliverables=[
                "Site visit and inspection",
                "Construction progress documentation",
                "Compliance verification",
                "Photo documentation",
                "Field observation report",
                "Punch list development",
            ],
            unit="per visit"
        ),
        "STORMWATER": ServicePricing(
            service_id="STORMWATER",
            service_name="Stormwater Management Plan (SWMP)",
            description="Comprehensive stormwater pollution prevention and management planning",
            base_price=Decimal("3800"),
            time_estimate_days=9,
            deliverables=[
                "SWMP document preparation",
                "Best Management Practices (BMP) design",
                "Erosion and sediment control plan",
                "LPDES compliance documentation",
                "Inspection and maintenance procedures",
                "SWMP implementation guidance",
            ],
            unit="per project"
        ),
    }

    def calculate_total(
        self,
        services: List[str],
        discount_percent: Decimal = Decimal("0"),
        rush_fee_percent: Decimal = Decimal("0")
    ) -> Dict:
        """
        Calculate total pricing for selected civil engineering services.

        Args:
            services: List of service IDs (e.g., ["DIA", "GRADING", "REVIEW"])
            discount_percent: Custom discount percentage (0-100)
            rush_fee_percent: Rush fee percentage (0-100)

        Returns:
            Dictionary with pricing breakdown
        """
        # Get service details
        selected_services = [
            self.SERVICE_PRICING[s]
            for s in services
            if s in self.SERVICE_PRICING
        ]

        # Calculate subtotal
        subtotal = sum(s.base_price for s in selected_services)

        # Apply package discount
        package_discount = self._calculate_package_discount(len(services))
        total_discount = max(discount_percent, package_discount)

        # Calculate discount amount
        discount_amount = subtotal * (total_discount / Decimal("100"))

        # Subtotal after discount
        discounted_subtotal = subtotal - discount_amount

        # Apply rush fee if applicable
        rush_fee_amount = discounted_subtotal * (rush_fee_percent / Decimal("100"))

        # Calculate total
        total = discounted_subtotal + rush_fee_amount

        # Estimate timeline
        total_days = sum(s.time_estimate_days for s in selected_services)

        return {
            "services": selected_services,
            "subtotal": float(subtotal),
            "package_discount_percent": float(package_discount),
            "custom_discount_percent": float(discount_percent),
            "total_discount_percent": float(total_discount),
            "discount_amount": float(discount_amount),
            "discounted_subtotal": float(discounted_subtotal),
            "rush_fee_percent": float(rush_fee_percent),
            "rush_fee_amount": float(rush_fee_amount),
            "total": float(total),
            "estimated_days": total_days,
            "estimated_completion": (datetime.now() + timedelta(days=total_days)).strftime("%B %d, %Y"),
        }

    def _calculate_package_discount(self, service_count: int) -> Decimal:
        """
        Calculate package discount based on number of services.

        Args:
            service_count: Number of services selected

        Returns:
            Discount percentage
        """
        if service_count >= 7:
            return Decimal("15")  # 15% for comprehensive package
        elif service_count >= 5:
            return Decimal("10")  # 10% for 5+ services
        elif service_count >= 3:
            return Decimal("5")   # 5% for 3+ services
        else:
            return Decimal("0")   # No package discount


class ProposalGenerator:
    """
    Generate professional branded proposals for civil engineering services.

    Creates Word documents with:
    - Cover page with LCR & Company branding
    - Project overview and scope
    - Detailed service descriptions
    - Pricing breakdown
    - Timeline and milestones
    - Terms and conditions
    - Signature page
    """

    def __init__(self, output_dir: str = "/app/outputs"):
        """
        Initialize proposal generator.

        Args:
            output_dir: Directory for output files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.calculator = PricingCalculator()

    def generate_proposal(
        self,
        proposal_data: ProposalData,
        output_filename: Optional[str] = None
    ) -> str:
        """
        Generate complete civil engineering services proposal document.

        Args:
            proposal_data: Proposal data and requirements
            output_filename: Optional custom filename

        Returns:
            Path to generated proposal file
        """
        logger.info(f"Generating proposal for {proposal_data.client_name}")

        # Calculate pricing
        pricing = self.calculator.calculate_total(
            services=proposal_data.services_requested,
            discount_percent=proposal_data.discount_percent,
            rush_fee_percent=proposal_data.rush_fee_percent
        )

        # Create document
        doc = Document()

        # Cover page
        self._add_cover_page(doc, proposal_data)
        doc.add_page_break()

        # Project overview
        self._add_project_overview(doc, proposal_data)
        doc.add_page_break()

        # Scope of services
        self._add_scope_of_services(doc, proposal_data, pricing)
        doc.add_page_break()

        # Pricing
        self._add_pricing_section(doc, pricing)
        doc.add_page_break()

        # Timeline
        self._add_timeline_section(doc, pricing)
        doc.add_page_break()

        # Terms and conditions
        self._add_terms_and_conditions(doc, proposal_data)
        doc.add_page_break()

        # Signature page
        self._add_signature_page(doc, proposal_data)

        # Save document
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d")
            safe_client = proposal_data.client_name.replace(" ", "_")
            output_filename = f"Proposal_{safe_client}_{timestamp}.docx"

        output_path = self.output_dir / output_filename
        doc.save(str(output_path))

        logger.info(f"Generated proposal: {output_path}")
        return str(output_path)

    def _add_cover_page(self, doc: Document, data: ProposalData):
        """Add branded cover page with LCR branding"""
        # Company logo/header (would add actual logo in production)
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        company_name = "LCR & COMPANY"
        tagline = "Civil Engineering & Land Surveying Services"

        run = header.add_run(f"{company_name}\\n")
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        run = header.add_run(f"{tagline}\\n\\n\\n")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Proposal title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("PROFESSIONAL SERVICES PROPOSAL\\n\\n\\n")
        run.font.size = Pt(20)
        run.font.bold = True

        # Project info
        project_info = doc.add_paragraph()
        project_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = project_info.add_run(f"{data.project_name}\\n")
        run.font.size = Pt(18)
        run.font.bold = True

        run = project_info.add_run(f"{data.project_location}\\n\\n\\n")
        run.font.size = Pt(14)

        # Client info
        client_info = doc.add_paragraph()
        client_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = client_info.add_run("Prepared For:\\n")
        run.font.size = Pt(12)
        run.font.bold = True

        run = client_info.add_run(
            f"{data.client_name}\\n"
            f"{data.client_contact}\\n"
            f"{data.client_email}\\n\\n\\n"
        )
        run.font.size = Pt(12)

        # Date and validity
        date_info = doc.add_paragraph()
        date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        proposal_date = datetime.now().strftime("%B %d, %Y")
        valid_until = (datetime.now() + timedelta(days=data.proposal_valid_days)).strftime("%B %d, %Y")

        run = date_info.add_run(
            f"Proposal Date: {proposal_date}\\n"
            f"Valid Until: {valid_until}\\n\\n"
        )
        run.font.size = Pt(11)

        # Prepared by
        doc.add_paragraph("\\n" * 5)
        prepared = doc.add_paragraph()
        prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = prepared.add_run(f"Prepared By:\\n{data.prepared_by}")
        run.font.size = Pt(11)

    def _add_project_overview(self, doc: Document, data: ProposalData):
        """Add project overview section"""
        doc.add_heading("PROJECT OVERVIEW", level=1)

        overview = doc.add_paragraph()
        overview.add_run("Project Name: ").font.bold = True
        overview.add_run(f"{data.project_name}\\n\\n")

        overview.add_run("Location: ").font.bold = True
        overview.add_run(f"{data.project_location}\\n\\n")

        overview.add_run("Project Type: ").font.bold = True
        overview.add_run(f"{data.project_type}\\n\\n")

        overview.add_run("Jurisdiction: ").font.bold = True
        overview.add_run(f"{data.jurisdiction}\\n\\n")

        overview.add_run("Description:\\n").font.bold = True
        overview.add_run(f"{data.project_description}\\n\\n")

        # About section
        doc.add_heading("ABOUT LCR & COMPANY", level=2)

        about_text = (
            "LCR & Company is a professional civil engineering and land surveying firm "
            "specializing in drainage analysis, stormwater management, and site development services. "
            "Our team of licensed professional engineers has extensive experience with Lafayette UDC, "
            "DOTD, and LPDES regulatory requirements.\\n\\n"
            "We provide comprehensive civil engineering services including drainage impact analysis, "
            "grading design, detention pond design, plan review, and construction observation. "
            "All deliverables are prepared in accordance with local, state, and federal standards, "
            "ensuring regulatory compliance and timely project approvals."
        )
        doc.add_paragraph(about_text)

    def _add_scope_of_services(
        self,
        doc: Document,
        data: ProposalData,
        pricing: Dict
    ):
        """Add scope of services section"""
        doc.add_heading("SCOPE OF SERVICES", level=1)

        doc.add_paragraph(
            "LCR & Company proposes to provide the following civil engineering services "
            "for this project:"
        )

        # Add each selected service
        for service in pricing["services"]:
            doc.add_heading(f"{service.service_name}", level=2)

            desc_para = doc.add_paragraph()
            desc_para.add_run("Service Description: ").font.bold = True
            desc_para.add_run(f"{service.description}\\n\\n")

            desc_para.add_run("Deliverables:\\n").font.bold = True
            for deliverable in service.deliverables:
                doc.add_paragraph(deliverable, style='List Bullet')

            # Add unit pricing info if applicable
            unit_para = doc.add_paragraph()
            unit_para.add_run(f"\\nPricing: ").font.bold = True
            unit_para.add_run(f"${float(service.base_price):,.2f} {service.unit}")

            doc.add_paragraph("")  # Spacing

        # Custom scope items
        if data.custom_scope:
            doc.add_heading("Additional Custom Services", level=2)
            for item in data.custom_scope:
                doc.add_paragraph(item, style='List Bullet')

    def _add_pricing_section(self, doc: Document, pricing: Dict):
        """Add pricing breakdown section"""
        doc.add_heading("PROFESSIONAL FEE ESTIMATE", level=1)

        # Create pricing table
        table = doc.add_table(rows=len(pricing["services"]) + 6, cols=3)
        table.style = 'Light Grid Accent 1'

        # Header
        headers = ["Service", "Description", "Fee"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Service rows
        for idx, service in enumerate(pricing["services"], start=1):
            row = table.rows[idx]
            row.cells[0].text = service.service_id
            row.cells[1].text = service.service_name
            row.cells[2].text = f"${service.base_price:,.2f}"

        # Subtotal
        row_idx = len(pricing["services"]) + 1
        table.rows[row_idx].cells[1].text = "Subtotal"
        table.rows[row_idx].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[row_idx].cells[2].text = f"${pricing['subtotal']:,.2f}"

        # Discount
        if pricing["total_discount_percent"] > 0:
            row_idx += 1
            discount_label = f"Discount ({pricing['total_discount_percent']:.0f}%)"
            if pricing["package_discount_percent"] > 0:
                discount_label += " - Package Discount!"
            table.rows[row_idx].cells[1].text = discount_label
            table.rows[row_idx].cells[2].text = f"-${pricing['discount_amount']:,.2f}"
            table.rows[row_idx].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

        # Rush fee
        if pricing["rush_fee_percent"] > 0:
            row_idx += 1
            table.rows[row_idx].cells[1].text = f"Rush Fee ({pricing['rush_fee_percent']:.0f}%)"
            table.rows[row_idx].cells[2].text = f"+${pricing['rush_fee_amount']:,.2f}"

        # Total
        row_idx += 1
        table.rows[row_idx].cells[1].text = "TOTAL INVESTMENT"
        table.rows[row_idx].cells[1].paragraphs[0].runs[0].font.bold = True
        table.rows[row_idx].cells[1].paragraphs[0].runs[0].font.size = Pt(14)
        table.rows[row_idx].cells[2].text = f"${pricing['total']:,.2f}"
        table.rows[row_idx].cells[2].paragraphs[0].runs[0].font.bold = True
        table.rows[row_idx].cells[2].paragraphs[0].runs[0].font.size = Pt(14)

        doc.add_paragraph("")
        payment_terms = doc.add_paragraph()
        payment_terms.add_run("Payment Terms:\\n").font.bold = True
        payment_terms.add_run(
            "• 50% due upon contract execution\\n"
            "• 50% due upon delivery of final deliverables\\n"
            "• Payment accepted via check or ACH transfer\\n"
        )

    def _add_timeline_section(self, doc: Document, pricing: Dict):
        """Add timeline and milestones section"""
        doc.add_heading("PROJECT TIMELINE", level=1)

        timeline_para = doc.add_paragraph()
        timeline_para.add_run("Estimated Duration: ").font.bold = True
        timeline_para.add_run(f"{pricing['estimated_days']} business days\\n\\n")

        timeline_para.add_run("Estimated Completion: ").font.bold = True
        timeline_para.add_run(f"{pricing['estimated_completion']}\\n\\n")

        timeline_para.add_run("Note: ").font.bold = True
        timeline_para.add_run(
            "Timeline assumes timely receipt of all necessary data, survey information, "
            "and client approvals. Delays in data provision or agency review may extend the timeline.\\n\\n"
        )

        doc.add_heading("Project Milestones", level=2)

        milestones = [
            "Contract Execution & Project Kickoff",
            "Data Collection (plans, survey, site information)",
            "Technical Analysis & Calculations",
            "Draft Report/Plans Preparation",
            "Internal QA Review",
            "Client Review & Comment Period",
            "Revisions & Final Document Preparation",
            "Final Deliverable Submittal",
        ]

        for milestone in milestones:
            doc.add_paragraph(milestone, style='List Bullet')

    def _add_terms_and_conditions(self, doc: Document, data: ProposalData):
        """Add terms and conditions"""
        doc.add_heading("TERMS & CONDITIONS", level=1)

        terms = [
            ("Scope Changes", "Any changes to the scope of work will be documented via written change order with associated cost and schedule adjustments."),
            ("Client Responsibilities", "Client will provide timely access to project data, site plans, survey data, and necessary regulatory documents. Client is responsible for obtaining all required permits and approvals."),
            ("Professional Standards", "All services will be performed in accordance with applicable professional standards of care for civil engineering practice in Louisiana."),
            ("Regulatory Compliance", "All deliverables will be prepared to meet Lafayette UDC, DOTD, LPDES, and other applicable regulatory requirements as specified."),
            ("Ownership of Documents", "All reports, plans, and technical documents prepared by LCR & Company become the property of the client upon receipt of final payment."),
            ("Insurance & Liability", "LCR & Company maintains professional liability insurance. Liability is limited to the fee paid for services rendered."),
            ("Confidentiality", "All project data and deliverables will be kept confidential and used only for this project unless otherwise agreed."),
            ("Dispute Resolution", "Any disputes arising from this agreement will be resolved through mediation or arbitration in Lafayette Parish, Louisiana."),
        ]

        for title, description in terms:
            term_para = doc.add_paragraph()
            term_para.add_run(f"{title}: ").font.bold = True
            term_para.add_run(f"{description}\\n")

    def _add_signature_page(self, doc: Document, data: ProposalData):
        """Add signature page"""
        doc.add_heading("ACCEPTANCE", level=1)

        acceptance = doc.add_paragraph(
            "By signing below, client accepts the scope, pricing, and terms outlined in this proposal "
            "and authorizes LCR & Company to proceed with the described services."
        )

        doc.add_paragraph("\\n" * 2)

        # Client signature block
        client_sig = doc.add_paragraph()
        client_sig.add_run("CLIENT ACCEPTANCE:\\n\\n")
        client_sig.add_run("_" * 50 + "\\n")
        client_sig.add_run(f"{data.client_contact}, {data.client_name}\\n\\n")
        client_sig.add_run("Date: _" + "_" * 30 + "\\n\\n\\n")

        # Company signature block
        company_sig = doc.add_paragraph()
        company_name = "LCR & COMPANY"

        company_sig.add_run(f"{company_name.upper()}:\\n\\n")
        company_sig.add_run("_" * 50 + "\\n")
        company_sig.add_run(f"{data.prepared_by}\\n\\n")
        company_sig.add_run("Date: _" + "_" * 30 + "\\n")
