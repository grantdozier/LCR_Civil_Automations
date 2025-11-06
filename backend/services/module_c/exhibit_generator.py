"""
Module C - Exhibit Generator
Generate technical exhibits (3A-3D) for DIA reports
"""
from typing import Dict, List, Optional
from pathlib import Path
from datetime import datetime
import logging

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class ExhibitGenerator:
    """
    Generate technical exhibits for DIA reports.

    Exhibits include:
    - Exhibit 3A: 10-Year Storm Analysis
    - Exhibit 3B: 25-Year Storm Analysis
    - Exhibit 3C: 50-Year Storm Analysis
    - Exhibit 3D: 100-Year Storm Analysis

    Each exhibit contains:
    - Drainage area summary table
    - Time of Concentration calculations
    - Rational Method calculations
    - Peak flow results
    """

    def __init__(self, output_dir: str = "/app/outputs"):
        """
        Initialize exhibit generator.

        Args:
            output_dir: Directory for output files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_all_exhibits(
        self,
        project_data: Dict,
        drainage_areas: List[Dict],
        results: Dict[str, List[Dict]],
        output_prefix: str = "Exhibit"
    ) -> List[str]:
        """
        Generate all exhibits (3A-3D).

        Args:
            project_data: Project information
            drainage_areas: List of drainage area data
            results: Dictionary mapping storm event to results
            output_prefix: Prefix for output filenames

        Returns:
            List of paths to generated exhibit files
        """
        exhibit_files = []

        storm_exhibits = {
            "3A": "10-year",
            "3B": "25-year",
            "3C": "50-year",
            "3D": "100-year",
        }

        for exhibit_id, storm_event in storm_exhibits.items():
            if storm_event in results:
                exhibit_path = self.generate_exhibit(
                    exhibit_id=exhibit_id,
                    storm_event=storm_event,
                    project_data=project_data,
                    drainage_areas=drainage_areas,
                    storm_results=results[storm_event],
                    output_prefix=output_prefix
                )
                exhibit_files.append(exhibit_path)

        logger.info(f"Generated {len(exhibit_files)} exhibits")
        return exhibit_files

    def generate_exhibit(
        self,
        exhibit_id: str,
        storm_event: str,
        project_data: Dict,
        drainage_areas: List[Dict],
        storm_results: List[Dict],
        output_prefix: str = "Exhibit"
    ) -> str:
        """
        Generate single exhibit document.

        Args:
            exhibit_id: Exhibit identifier (e.g., "3A", "3B")
            storm_event: Storm event (e.g., "10-year", "25-year")
            project_data: Project information
            drainage_areas: List of drainage area data
            storm_results: Results for this storm event
            output_prefix: Prefix for output filename

        Returns:
            Path to generated exhibit file
        """
        logger.info(f"Generating Exhibit {exhibit_id} - {storm_event} Storm")

        # Create document
        doc = Document()

        # Title page
        self._add_exhibit_title(doc, exhibit_id, storm_event, project_data)
        doc.add_page_break()

        # Drainage area summary
        self._add_drainage_area_summary(doc, drainage_areas)
        doc.add_page_break()

        # Time of Concentration calculations
        self._add_tc_calculations(doc, storm_results)
        doc.add_page_break()

        # Rational Method calculations
        self._add_rational_method_table(doc, storm_event, storm_results)
        doc.add_page_break()

        # Results summary
        self._add_results_summary(doc, storm_event, storm_results)

        # Save document
        timestamp = datetime.now().strftime("%Y%m%d")
        filename = f"{output_prefix}_{exhibit_id}_{storm_event}_{timestamp}.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))

        logger.info(f"Generated exhibit: {output_path}")
        return str(output_path)

    def _add_exhibit_title(self, doc: Document, exhibit_id: str, storm_event: str, project_data: Dict):
        """Add exhibit title page"""
        # Main title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"EXHIBIT {exhibit_id}\n\n")
        run.font.size = Pt(24)
        run.font.bold = True

        # Storm event
        storm_title = doc.add_paragraph()
        storm_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = storm_title.add_run(f"{storm_event.upper()} STORM EVENT\n")
        run.font.size = Pt(20)
        run.font.bold = True

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("DRAINAGE ANALYSIS CALCULATIONS\n\n\n")
        run.font.size = Pt(14)

        # Project info
        project_info = doc.add_paragraph()
        project_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = project_info.add_run(f"{project_data.get('project_name', '')}\n")
        run.font.size = Pt(16)
        run.font.bold = True

        location = doc.add_paragraph()
        location.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = location.add_run(f"{project_data.get('location', '')}\n\n")
        run.font.size = Pt(12)

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(project_data.get('date', datetime.now().strftime('%B %d, %Y')))
        run.font.size = Pt(12)

    def _add_drainage_area_summary(self, doc: Document, drainage_areas: List[Dict]):
        """Add drainage area summary table"""
        doc.add_heading("DRAINAGE AREA SUMMARY", level=1)

        # Create table
        table = doc.add_table(rows=len(drainage_areas) + 1, cols=7)
        table.style = 'Light Grid Accent 1'

        # Headers
        headers = [
            "Area ID",
            "Total Area\n(acres)",
            "Impervious\n(acres)",
            "Pervious\n(acres)",
            "Impervious\n(%)",
            "Weighted C",
            "Notes"
        ]

        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for idx, da in enumerate(drainage_areas, start=1):
            row = table.rows[idx]
            row.cells[0].text = da.get('area_label', '')
            row.cells[1].text = f"{da.get('total_area_acres', 0):.2f}"
            row.cells[2].text = f"{da.get('impervious_area_acres', 0):.2f}"
            row.cells[3].text = f"{da.get('pervious_area_acres', 0):.2f}"
            row.cells[4].text = f"{da.get('impervious_percentage', 0):.1f}%"
            row.cells[5].text = f"{da.get('weighted_c_value', 0):.3f}"
            row.cells[6].text = da.get('notes', '')

            # Center align numeric cells
            for cell_idx in [1, 2, 3, 4, 5]:
                row.cells[cell_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_tc_calculations(self, doc: Document, storm_results: List[Dict]):
        """Add Time of Concentration calculations table"""
        doc.add_heading("TIME OF CONCENTRATION CALCULATIONS", level=1)

        intro = doc.add_paragraph(
            "Time of Concentration (Tc) is the time required for water to travel from the "
            "hydraulically most distant point to the outlet. The Tc is used to determine the "
            "appropriate rainfall intensity for each drainage area."
        )

        # Create table
        table = doc.add_table(rows=len(storm_results) + 1, cols=6)
        table.style = 'Light Grid Accent 1'

        # Headers
        headers = [
            "Area ID",
            "Flow Length\n(feet)",
            "Elevation\nChange (ft)",
            "Slope\n(%)",
            "Tc\n(minutes)",
            "Method"
        ]

        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for idx, result in enumerate(storm_results, start=1):
            row = table.rows[idx]
            row.cells[0].text = result.get('area_label', '')
            row.cells[1].text = str(result.get('flow_length_ft', '-'))
            row.cells[2].text = str(result.get('elevation_change_ft', '-'))
            row.cells[3].text = str(result.get('slope_percent', '-'))
            row.cells[4].text = f"{result.get('tc_minutes', 0):.2f}"
            row.cells[5].text = result.get('tc_method', 'NRCS')

            # Center align
            for cell_idx in range(6):
                row.cells[cell_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_rational_method_table(self, doc: Document, storm_event: str, storm_results: List[Dict]):
        """Add Rational Method calculations table"""
        doc.add_heading(f"RATIONAL METHOD CALCULATIONS - {storm_event.upper()}", level=1)

        # Formula explanation
        formula = doc.add_paragraph()
        formula.add_run("Rational Method Formula: Q = CiA\n\n").font.bold = True
        formula.add_run("Where:\n")
        formula.add_run("  Q = Peak runoff rate (cubic feet per second, cfs)\n")
        formula.add_run("  C = Weighted runoff coefficient (dimensionless)\n")
        formula.add_run("  i = Rainfall intensity (inches per hour)\n")
        formula.add_run("  A = Drainage area (acres)\n\n")

        # Create table
        table = doc.add_table(rows=len(storm_results) + 1, cols=7)
        table.style = 'Light Grid Accent 1'

        # Headers
        headers = [
            "Area ID",
            "C\n(coefficient)",
            "i\n(in/hr)",
            "A\n(acres)",
            "Tc\n(min)",
            "Q\n(cfs)",
            "Condition"
        ]

        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for idx, result in enumerate(storm_results, start=1):
            row = table.rows[idx]
            row.cells[0].text = result.get('area_label', '')
            row.cells[1].text = f"{result.get('c_value', 0):.3f}"
            row.cells[2].text = f"{result.get('i_value', 0):.2f}"
            row.cells[3].text = f"{result.get('area_acres', 0):.2f}"
            row.cells[4].text = f"{result.get('tc_minutes', 0):.2f}"
            row.cells[5].text = f"{result.get('peak_flow_cfs', 0):.1f}"
            row.cells[6].text = result.get('development_condition', 'post').title()

            # Center align
            for cell_idx in range(7):
                row.cells[cell_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add totals row
        total_row = table.add_row()
        total_row.cells[0].text = "TOTAL"
        total_row.cells[0].paragraphs[0].runs[0].font.bold = True
        total_flow = sum(r.get('peak_flow_cfs', 0) for r in storm_results)
        total_row.cells[5].text = f"{total_flow:.1f}"
        total_row.cells[5].paragraphs[0].runs[0].font.bold = True
        total_row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_results_summary(self, doc: Document, storm_event: str, storm_results: List[Dict]):
        """Add results summary and notes"""
        doc.add_heading("SUMMARY AND NOTES", level=1)

        # Summary statistics
        total_area = sum(r.get('area_acres', 0) for r in storm_results)
        total_flow = sum(r.get('peak_flow_cfs', 0) for r in storm_results)
        avg_c = sum(r.get('c_value', 0) for r in storm_results) / len(storm_results) if storm_results else 0

        summary = doc.add_paragraph()
        summary.add_run("Analysis Summary:\n\n").font.bold = True
        summary.add_run(f"Storm Event: {storm_event.upper()}\n")
        summary.add_run(f"Total Drainage Area: {total_area:.2f} acres\n")
        summary.add_run(f"Total Peak Flow: {total_flow:.1f} cfs\n")
        summary.add_run(f"Average Runoff Coefficient: {avg_c:.3f}\n")
        summary.add_run(f"Number of Drainage Areas: {len(storm_results)}\n\n")

        # Notes
        notes = doc.add_paragraph()
        notes.add_run("Notes:\n\n").font.bold = True
        notes.add_run("1. Rainfall intensity data obtained from NOAA Atlas 14, Volume 9.\n")
        notes.add_run("2. Time of Concentration calculated using NRCS method.\n")
        notes.add_run("3. Runoff coefficients based on Lafayette UDC and site conditions.\n")
        notes.add_run("4. All calculations performed in accordance with accepted engineering practice.\n")
        notes.add_run("5. Post-development conditions reflect proposed site improvements.\n")

    def generate_noaa_appendix(
        self,
        project_data: Dict,
        intensities: Dict[str, Dict[float, float]]
    ) -> str:
        """
        Generate NOAA Atlas 14 data appendix.

        Args:
            project_data: Project information
            intensities: Dictionary mapping storm event to {duration: intensity}
                       Example: {
                           "10-year": {5: 8.92, 10: 7.25, 15: 6.38, 30: 4.85},
                           "25-year": {5: 10.65, 10: 8.65, 15: 7.62, 30: 5.79}
                       }

        Returns:
            Path to generated appendix file
        """
        logger.info("Generating NOAA Atlas 14 Appendix")

        doc = Document()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("APPENDIX A\n\n")
        run.font.size = Pt(24)
        run.font.bold = True

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("NOAA ATLAS 14 RAINFALL INTENSITY DATA\n\n")
        run.font.size = Pt(18)
        run.font.bold = True

        # Location info
        doc.add_paragraph(
            f"Location: {project_data.get('location', 'Lafayette, Louisiana')}\n"
            "Data Source: NOAA Atlas 14, Volume 9 (Southeastern States)\n\n"
        )

        # Create intensity table
        doc.add_heading("Rainfall Intensity (inches per hour)", level=2)

        # Get all unique durations
        all_durations = set()
        for storm_data in intensities.values():
            all_durations.update(storm_data.keys())
        durations = sorted(all_durations)

        # Create table
        table = doc.add_table(rows=len(intensities) + 1, cols=len(durations) + 1)
        table.style = 'Light Grid Accent 1'

        # Headers - durations
        table.rows[0].cells[0].text = "Storm Event"
        table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
        for idx, duration in enumerate(durations, start=1):
            cell = table.rows[0].cells[idx]
            cell.text = f"{int(duration)} min"
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for row_idx, (storm_event, storm_data) in enumerate(intensities.items(), start=1):
            table.rows[row_idx].cells[0].text = storm_event.title()
            table.rows[row_idx].cells[0].paragraphs[0].runs[0].font.bold = True

            for col_idx, duration in enumerate(durations, start=1):
                intensity = storm_data.get(duration, 0)
                table.rows[row_idx].cells[col_idx].text = f"{intensity:.2f}"
                table.rows[row_idx].cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save
        filename = f"Appendix_A_NOAA_Atlas_14_{datetime.now().strftime('%Y%m%d')}.docx"
        output_path = self.output_dir / filename
        doc.save(str(output_path))

        logger.info(f"Generated NOAA appendix: {output_path}")
        return str(output_path)
